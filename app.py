import streamlit as st
from io import BytesIO, StringIO
import requests
import pandas as pd
from datetime import datetime, timedelta, timezone
from openpyxl import load_workbook
from openpyxl.styles import Alignment
import shutil
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # Needed for setting cell shading

# ---- Page Configuration ----
st.set_page_config(
    page_title="Operations - Weekly Training Plan",
    layout="wide",  # Use wide layout
    initial_sidebar_state="expanded"
)

# ---- Custom CSS to hide default Streamlit elements and reduce top spacing ----
hide_streamlit_style = """
<style>
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# ----------------------------------------
# Helper function to set cell background color in Word tables
def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

# ----------------------------------------
# Function to adjust timestamps and convert to local time
def convert_to_time(timestamp_ms, offset_hours=11):
    try:
        if pd.notnull(timestamp_ms):
            timestamp_s = float(timestamp_ms) / 1000
            return (datetime.fromtimestamp(timestamp_s, tz=timezone.utc) - timedelta(hours=offset_hours)).strftime('%H:%M')
    except (ValueError, TypeError):
        return None

# ----------------------------------------
# Function to ensure all expected columns are present in the pivot DataFrame
def ensure_all_columns(pivot_df, day_order):
    return pivot_df.reindex(columns=['Sport', 'Training_Group'] + day_order, fill_value=' ')

# ----------------------------------------
# Updated function to format session information for grouping in the Excel calendar.
def format_session(group):
    tc_sessions = []    # For training camp rows
    other_sessions = [] # For other sessions
    for _, row in group.iterrows():
        session_type = str(row.get('Session_Type', '')).strip().lower()
        if session_type == "training camp":
            st_time = str(row['Start_Time']) if pd.notnull(row['Start_Time']) and str(row['Start_Time']).strip() != "" else "23:59"
            tc_sessions.append((st_time, "TRAINING CAMP"))
        else:
            venue = str(row['Venue']) if pd.notnull(row['Venue']) else ''
            start_time = str(row['Start_Time']) if pd.notnull(row['Start_Time']) else ''
            finish_time = str(row['Finish_Time']) if pd.notnull(row['Finish_Time']) else ''
            time_info = f"{start_time}-{finish_time}" if start_time or finish_time else ''
            detail = f"{venue}\n{time_info}".strip()
            other_sessions.append((start_time, detail))
    def sort_key(x):
        try:
            return datetime.strptime(x[0], '%H:%M')
        except Exception:
            return datetime.max
    tc_sessions_sorted = sorted(tc_sessions, key=sort_key)
    other_sessions_sorted = sorted(other_sessions, key=sort_key)
    if tc_sessions_sorted and other_sessions_sorted:
        tc_str = " & ".join(s[1] for s in tc_sessions_sorted if s[1])
        other_str = "\n".join(s[1] for s in other_sessions_sorted if s[1])
        return f"{tc_str}\n{other_str}"
    else:
        combined = tc_sessions_sorted + other_sessions_sorted
        combined_sorted = sorted(combined, key=sort_key)
        return "\n".join(s[1] for s in combined_sorted if s[1])

# ----------------------------------------
# Function to paste filtered data into the Excel template sheet
def paste_filtered_data_to_template(pivot_df, workbook, sport, training_group, start_cell):
    template_sheet = workbook["Template"]
    filtered_row = pivot_df[(pivot_df['Sport'] == sport) & (pivot_df['Training_Group'] == training_group)]
    if not filtered_row.empty:
        values_to_paste = filtered_row.iloc[0, 2:].tolist()  # Exclude Sport and Training_Group columns
        col_letter, row_num = start_cell[0], int(start_cell[1:])
        start_col_idx = ord(col_letter.upper()) - ord("A") + 1
        for col_idx, value in enumerate(values_to_paste, start=start_col_idx):
            cell = template_sheet.cell(row=row_num, column=col_idx, value=value)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

# ----------------------------------------
# Function to generate the Excel report (Calendar)
def generate_excel(selected_date):
    template_path = "Excel_template.xlsx"  # Path to your Excel template
    output_filename = f"Training_Report_{selected_date.strftime('%d%b%Y')}.xlsx"
    if os.path.exists(output_filename):
        os.remove(output_filename)
    shutil.copy(template_path, output_filename)
    workbook = load_workbook(output_filename)
    template_sheet = workbook["Template"]

    start_date = selected_date
    end_date = start_date + timedelta(days=6)
    st.write(f"**Selected Date Range:** {start_date.strftime('%a %d %b %Y')} to {end_date.strftime('%a %d %b %Y')}")

    session_req = requests.Session()
    session_req.auth = ("sb_sap.etl", "A1s2p3!re")
    response = session_req.get("https://aspire.smartabase.com/aspireacademy/live?report=PYTHON6_TRAINING_PLAN&updategroup=true")
    response.raise_for_status()
    data = pd.read_html(StringIO(response.text))[0]
    df = data.drop(columns=['About'], errors='ignore').drop_duplicates()
    df.columns = df.columns.str.replace(' ', '_')

    # Normalize sport and training group values
    df['Sport'] = df['Sport'].str.replace(' ', '_')
    df['Training_Group'] = df['Training_Group'].str.replace(' ', '_')
    
    df['Start_Time'] = pd.to_numeric(df['Start_Time'], errors='coerce').apply(lambda x: convert_to_time(x))
    df['Finish_Time'] = pd.to_numeric(df['Finish_Time'], errors='coerce').apply(lambda x: convert_to_time(x))
    
    df = df[df['Sport'].notna() & (df['Sport'].str.strip() != '')]
    df = df[df['Venue'] != 'AASMC']
    df = df[df['Sport'] != 'Generic_Athlete']
    df = df[df['Training_Group'] != 'Practice']

    df['Date'] = pd.to_datetime(df['Date'], errors='coerce', dayfirst=True).dt.date
    filtered_df = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
    if 'AM/PM' in filtered_df.columns:
        filtered_df.loc[:, 'AM/PM'] = pd.Categorical(filtered_df['AM/PM'], categories=['AM', 'PM'], ordered=True)
    filtered_df = filtered_df.dropna(subset=['Sport']).sort_values(by=['Date', 'Sport', 'Coach', 'AM/PM'])

    calendar_grouped = filtered_df.groupby(['Sport', 'Training_Group', 'Day_AM/PM']).apply(format_session).reset_index()
    calendar_grouped.columns = ['Sport', 'Training_Group', 'Day_AM/PM', 'Session']

    pivot_df = pd.pivot_table(
        calendar_grouped,
        values='Session',
        index=['Sport', 'Training_Group'],
        columns=['Day_AM/PM'],
        aggfunc='first',
        fill_value=' '
    ).reset_index()

    day_order = [f"{day} {time}" for day in ['Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
                                       for time in ['AM', 'PM']]
    pivot_df = ensure_all_columns(pivot_df, day_order)

    # --- Updated Post-processing:
    # For each day, if one slot (AM or PM) contains "TRAINING CAMP" (case-insensitive),
    # then ensure both the AM and PM cells for that day include "TRAINING CAMP" as the first line.
    special_text = "TRAINING CAMP"
    def prepend_special(text):
        # Remove any leading occurrences of special_text (case-insensitive) and then prepend it.
        lines = text.splitlines()
        # Remove any line that is exactly the special text (case-insensitive).
        filtered = [line for line in lines if line.strip().upper() != special_text]
        return f"{special_text}" + ("\n" + "\n".join(filtered) if filtered else "")
    
    for day in ['Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']:
        am_col = f"{day} AM"
        pm_col = f"{day} PM"
        for idx, row in pivot_df.iterrows():
            am_val = str(row.get(am_col, '')).strip()
            pm_val = str(row.get(pm_col, '')).strip()
            am_has = special_text in am_val.upper()
            pm_has = special_text in pm_val.upper()
            if am_has or pm_has:
                # Prepend the special text to both cells.
                new_am = prepend_special(am_val)
                new_pm = prepend_special(pm_val)
                pivot_df.at[idx, am_col] = new_am
                pivot_df.at[idx, pm_col] = new_pm

    rows_to_paste = [
        {"sport": "Development", "training_group": "Development_1", "start_cell": "C6"},
        {"sport": "Development", "training_group": "Development_2", "start_cell": "C8"},
        {"sport": "Development", "training_group": "Development_3", "start_cell": "C10"},
        {"sport": "Endurance", "training_group": "Endurance_Senior", "start_cell": "C12"},
        {"sport": "Jumps", "training_group": "Jumps_PV", "start_cell": "C14"},
        {"sport": "Jumps", "training_group": "Jumps_Martin_Bercel", "start_cell": "C16"},
        {"sport": "Jumps", "training_group": "Jumps_Ross_Jeffs", "start_cell": "C18"},
        {"sport": "Jumps", "training_group": "Jumps_ElWalid", "start_cell": "C20"},
        {"sport": "Sprints", "training_group": "Sprints_Lee", "start_cell": "C22"},
        {"sport": "Sprints", "training_group": "Sprints_Hamdi", "start_cell": "C24"},
        {"sport": "Throws", "training_group": "Senior_Performance_Throws", "start_cell": "C26"},
        {"sport": "Squash", "training_group": "Squash", "start_cell": "C37"},
        {"sport": "Table_Tennis", "training_group": "Table_Tennis", "start_cell": "C39"},
        {"sport": "Fencing", "training_group": "Fencing", "start_cell": "C41"},
        {"sport": "Swimming", "training_group": "Swimming", "start_cell": "C43"},
        {"sport": "Padel", "training_group": "Padel", "start_cell": "C45"},
        {"sport": "Pre_Academy_Padel", "training_group": "Explorers", "start_cell": "C48"},
        {"sport": "Pre_Academy_Padel", "training_group": "Explorers+", "start_cell": "C49"},
        {"sport": "Pre_Academy_Padel", "training_group": "Starters", "start_cell": "C50"},
        {"sport": "Pre_Academy", "training_group": "Pre_Academy_Fencing", "start_cell": "C51"},
        {"sport": "Pre_Academy", "training_group": "Pre_Academy_Squash_Girls", "start_cell": "C53"},
        {"sport": "Pre_Academy", "training_group": "Pre_Academy_Athletics", "start_cell": "C55"},
        {"sport": "Girls_Programe", "training_group": "Kids", "start_cell": "C58"},
        {"sport": "Girls_Programe", "training_group": "Mini_Cadet_U14", "start_cell": "C59"},
        {"sport": "Girls_Programe", "training_group": "Cadet_U16", "start_cell": "C60"},
        {"sport": "Girls_Programe", "training_group": "Youth_U18", "start_cell": "C61"},
        {"sport": "Sprints", "training_group": "Sprints_Short", "start_cell": "C69"},
        {"sport": "Sprints", "training_group": "Sprints_Long", "start_cell": "C71"},
        {"sport": "Jumps", "training_group": "Jumps_QAF", "start_cell": "C73"}
    ]
    for row in rows_to_paste:
        paste_filtered_data_to_template(
            pivot_df=pivot_df,
            workbook=workbook,
            sport=row["sport"],
            training_group=row["training_group"],
            start_cell=row["start_cell"],
        )

    date_cells_groups = [
        ['C4', 'E4', 'G4', 'I4', 'K4', 'M4', 'O4'],
        ['C35', 'E35', 'G35', 'I35', 'K35', 'M35', 'O35'],
        ['C67', 'E67', 'G67', 'I67', 'K67', 'M67', 'O67'],
    ]
    for day_offset, cell_group in enumerate(zip(*date_cells_groups)):
        date_value = (start_date + timedelta(days=day_offset)).strftime('%a %d %b %Y')
        for cell in cell_group:
            template_sheet[cell].value = date_value
            template_sheet[cell].alignment = Alignment(horizontal="center", vertical="center")

    week_number = start_date.isocalendar()[1]
    week_beginning_text = f"Week beginning {start_date.strftime('%d %b')}\nWeek {week_number}"
    for cell in ["O2", "O33", "O65"]:
        template_sheet[cell].value = week_beginning_text
        template_sheet[cell].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    return output, pivot_df, filtered_df

# ----------------------------------------
# Function to generate a nicely formatted Word document (Venue Usage Report)
def generate_venue_usage_report(filtered_df, start_date):
    doc = Document()
    section = doc.sections[0]
    section.orientation = 1  # Landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    title = doc.add_heading('Venue Usage Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f'Week Beginning: {start_date.strftime("%d %b %Y")}', style='Normal')
    
    day_colors = {
        "Sunday": "D3D3D3",
        "Monday": "FFFFFF",
        "Tuesday": "D3D3D3",
        "Wednesday": "FFFFFF",
        "Thursday": "D3D3D3",
        "Friday": "FFFFFF",
        "Saturday": "D3D3D3"
    }
    
    venues = sorted([str(v) for v in filtered_df['Venue'].dropna().unique()])
    page_capacity = 5

    for i in range(0, len(venues), page_capacity):
        if i > 0:
            doc.add_page_break()
        venue_subset = venues[i:i+page_capacity]
        for venue in venue_subset:
            venue_data = filtered_df[filtered_df['Venue'].apply(lambda x: str(x)) == venue].sort_values(by=['Date', 'Start_Time'])
            venue_heading = doc.add_heading(f'📍 {venue}', level=2)
            venue_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Date'
            hdr_cells[1].text = 'Time'
            hdr_cells[2].text = 'Session Type'
            hdr_cells[3].text = 'Training Group'
            hdr_cells[4].text = 'Sport'
            for cell in hdr_cells:
                para = cell.paragraphs[0]
                para.runs[0].bold = True
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                set_cell_background(cell, "ADD8E6")
            
            for _, row in venue_data.iterrows():
                row_cells = table.add_row().cells
                date_str = row['Date'].strftime('%A %d %b %Y')
                time_str = f"{row['Start_Time']} - {row['Finish_Time']}"
                session_type = str(row['Session_Type']) if pd.notnull(row.get('Session_Type')) else ""
                row_cells[0].text = date_str
                row_cells[1].text = time_str
                row_cells[2].text = session_type
                row_cells[3].text = row['Training_Group']
                row_cells[4].text = row['Sport']
                
                day_name = row['Date'].strftime('%A')
                color = day_colors.get(day_name, "FFFFFF")
                for cell in row_cells:
                    set_cell_background(cell, color)
                    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

if "generated" not in st.session_state:
    st.session_state.generated = False
if "excel_file" not in st.session_state:
    st.session_state.excel_file = None
if "venue_file" not in st.session_state:
    st.session_state.venue_file = None
if "pivot_df" not in st.session_state:
    st.session_state.pivot_df = None
if "filtered_data" not in st.session_state:
    st.session_state.filtered_data = None

st.title("Operations - Weekly Training Plan App")
st.markdown("Generate Training Calendar and Venue Usage reports for any week from 1st January 2025.")

selected_date = st.date_input("Select a starting date (make sure to choose a SUNDAY!)", value=datetime.now().date())

if st.button("Generate Reports"):
    try:
        excel_file, pivot_df, filtered_data = generate_excel(selected_date)
        venue_file = generate_venue_usage_report(filtered_data, selected_date)
        st.session_state.excel_file = excel_file
        st.session_state.pivot_df = pivot_df
        st.session_state.filtered_data = filtered_data
        st.session_state.venue_file = venue_file
        st.session_state.generated = True
    except Exception as e:
        st.error(f"An error occurred: {e}")

if st.session_state.generated:
    st.markdown("### Pivot DataFrame for checking data")
    st.dataframe(st.session_state.pivot_df)
    st.download_button(
        label="📅 Download Training Calendar Excel Report",
        data=st.session_state.excel_file,
        file_name=f"Training_Report_{selected_date.strftime('%d%b%Y')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    st.download_button(
        label="📄 Download Venue Usage Report",
        data=st.session_state.venue_file,
        file_name=f"Venue_Usage_Report_{selected_date.strftime('%d%b%Y')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
