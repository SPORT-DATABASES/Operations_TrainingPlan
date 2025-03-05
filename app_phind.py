import streamlit as st
from io import BytesIO, StringIO 
import requests
import pandas as pd
from datetime import datetime, timedelta, timezone
from openpyxl import load_workbook
from openpyxl.styles import Alignment
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, List, Tuple, Optional
import logging
import shutil

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define constants
API_URL = "https://aspire.smartabase.com/aspireacademy/live"
REPORT_PARAMS = {"report": "PYTHON6_TRAINING_PLAN", "updategroup": True}


class TrainingReportGenerator:
    def __init__(self):
        self.session = requests.Session()
        self.session.auth = ("sb_sap.etl", "A1s2p3!re")
        
    @staticmethod
    def setup_page():
        """Configure Streamlit page settings."""
        st.set_page_config(
            page_title="Operations - Weekly Training Plan",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        st.markdown("""
            <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
            </style>
        """, unsafe_allow_html=True)

    @staticmethod
    def convert_timestamp(ms: float, offset_hours: int = 11) -> Optional[str]:
        """Convert millisecond timestamp to local time string."""
        try:
            if pd.isna(ms):
                return None
            dt = datetime.fromtimestamp(float(ms)/1000, tz=timezone.utc)
            return (dt - timedelta(hours=offset_hours)).strftime('%H:%M')
        except (ValueError, TypeError):
            logger.error(f"Failed to convert timestamp {ms}")
            return None

    def fetch_training_data(self) -> pd.DataFrame:
        """Fetch and preprocess training data."""
        try:
            logger.info("Fetching training data from API")
            response = self.session.get(API_URL, params=REPORT_PARAMS)
            response.raise_for_status()
            
            df = pd.read_html(StringIO(response.text))[0]
            
            # Clean and optimize dataframe
            if 'About' in df.columns:
                df.drop(columns=['About'], inplace=True)
            
            df.drop_duplicates(inplace=True)
            df.columns = df.columns.str.replace(' ', '_')
            
            # Convert numeric columns efficiently
            for col in ['Start_Time', 'Finish_Time']:
                df[col] = pd.to_numeric(df[col], errors='coerce')
                df[col] = df[col].apply(self.convert_timestamp)
            
            # Filter invalid entries with explicit type conversion
            mask = (
                df['Sport'].notna() & 
                df['Sport'].astype(str).str.strip().ne('') &
                df['Venue'].fillna('').ne('AASMC') &
                df['Sport'].fillna('').ne('Generic_Athlete') &
                df['Training_Group'].fillna('').ne('Practice')
            )
            
            logger.info(f"Fetched {len(df[mask])} valid training records")
            return df[mask].reset_index(drop=True)
            
        except Exception as e:
            logger.error(f"Failed to fetch training data: {str(e)}")
            raise

    def generate_excel_report(self, selected_date: datetime.date) -> Tuple[BytesIO, pd.DataFrame]:
        """Generate Excel report for the selected date range."""
        template_path = "Excel_template.xlsx"
        output_filename = f"Training_Report_{selected_date.strftime('%d%b%Y')}.xlsx"
        
        try:
            shutil.copy(template_path, output_filename)
            workbook = load_workbook(output_filename)
            template_sheet = workbook["Template"]
            
            df = self.fetch_training_data()
            df['Date'] = pd.to_datetime(df['Date'], errors='coerce', dayfirst=True).dt.date
            
            # Optimize filtering
            filtered_df = df[
                (df['Date'] >= selected_date) & 
                (df['Date'] <= selected_date + timedelta(days=6))
            ].copy()
            
            # Create efficient grouping
            pivot_df = (
                filtered_df.groupby(['Sport', 'Training_Group'])
                .agg({
                    'Session_Type': lambda x: '\n'.join(x),
                    'Venue': lambda x: '\n'.join(x),
                    'Start_Time': lambda x: '\n'.join(map(str, x)),
                    'Finish_Time': lambda x: '\n'.join(map(str, x))
                })
                .reset_index()
            )
            
            # Define sport configurations
            SPORT_CONFIGS = [
                {"sport": "Development", "group": "Development 1", "cell": "C6"},
                {"sport": "Development", "group": "Development 2", "cell": "C8"},
                {"sport": "Development", "training_group": "Development 3", "start_cell": "C10"},
        {"sport": "Endurance", "training_group": "Endurance_Senior", "start_cell": "C12"},
        {"sport": "Jumps", "training_group": "Jumps_PV", "start_cell": "C14"},
        {"sport": "Jumps", "training_group": "Jumps_Martin Bercel", "start_cell": "C16"},
        {"sport": "Jumps", "training_group": "Jumps_Ross Jeffs", "start_cell": "C18"},
        {"sport": "Jumps", "training_group": "Jumps_ElWalid", "start_cell": "C20"},
        {"sport": "Sprints", "training_group": "Sprints_Lee", "start_cell": "C22"},
        {"sport": "Sprints", "training_group": "Sprints_Hamdi", "start_cell": "C24"},
        {"sport": "Throws", "training_group": "Senior Performance Throws", "start_cell": "C26"},
        {"sport": "Squash", "training_group": "Squash", "start_cell": "C37"},
        {"sport": "Table Tennis", "training_group": "Table Tennis", "start_cell": "C39"},
        {"sport": "Fencing", "training_group": "Fencing", "start_cell": "C41"},
        {"sport": "Swimming", "training_group": "Swimming", "start_cell": "C43"},
        {"sport": "Padel", "training_group": "Padel", "start_cell": "C45"},
        {"sport": "Pre Academy Padel", "training_group": "Explorers", "start_cell": "C48"},
        {"sport": "Pre Academy Padel", "training_group": "Explorers+", "start_cell": "C49"},
        {"sport": "Pre Academy Padel", "training_group": "Starters", "start_cell": "C50"},
        {"sport": "Pre Academy", "training_group": "Pre Academy Fencing", "start_cell": "C51"},
        {"sport": "Pre Academy", "training_group": "Pre Academy Squash Girls", "start_cell": "C53"},
        {"sport": "Pre Academy", "training_group": "Pre Academy Athletics", "start_cell": "C55"},
        {"sport": "Girls Programe", "training_group": "Kids", "start_cell": "C58"},
        {"sport": "Girls Programe", "training_group": "Mini Cadet_U14", "start_cell": "C59"},
        {"sport": "Girls Programe", "training_group": "Cadet_U16", "start_cell": "C60"},
        {"sport": "Girls Programe", "training_group": "Youth_U18", "start_cell": "C61"},
        {"sport": "Sprints", "training_group": "Sprints_Short", "start_cell": "C69"},
        {"sport": "Sprints", "training_group": "Sprints_Long", "start_cell": "C71"},
        {"sport": "Jumps", "training_group": "Jumps_QAF", "start_cell": "C73"},
        {"sport": "Throws", "training_group": "Discus_QAF", "start_cell": "C75"},
        {"sport": "Throws", "training_group": "Hammer_QAF", "start_cell": "C77"},
        {"sport": "Throws", "training_group": "Javelin_QAF", "start_cell": "C79"},
      #  {"sport": "Throws", "training_group": "Javelin_QAF", "start_cell": "C79"},
       # {"sport": "Throws", "training_group": "Javelin_QAF", "start_cell": "C79"},
       # {"sport": "Throws", "training_group": "Javelin_QAF", "start_cell": "C79"},
       # {"sport": "Throws", "training_group": "Javelin_QAF", "start_cell": "C79"}
    ]
            
            
            # Batch update cells
            for config in SPORT_CONFIGS:
                filtered_row = pivot_df[
                    (pivot_df['Sport'] == config["sport"]) &
                    (pivot_df['Training_Group'] == config["group"])
                ].iloc[0].copy()
                
                col_letter, row_num = config["cell"][0], int(config["cell"][1:])
                start_col_idx = ord(col_letter.upper()) - ord("A") + 1
                
                for col_idx, value in enumerate(filtered_row.values[2:], start=start_col_idx):
                    cell = template_sheet.cell(row=row_num, column=col_idx, value=value)
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            
            # Update dates efficiently
            date_cells_groups = [
                   ['C4', 'E4', 'G4', 'I4', 'K4', 'M4', 'O4'],
        ['C35', 'E35', 'G35', 'I35', 'K35', 'M35', 'O35'],
        ['C67', 'E67', 'G67', 'I67', 'K67', 'M67', 'O67'],
            ]
            
            
            for day_offset, cell_group in enumerate(zip(*date_cells_groups)):
                date_value = (selected_date + timedelta(days=day_offset)).strftime('%a %d %b %Y')
                for cell in cell_group:
                    template_sheet[cell].value = date_value
                    template_sheet[cell].alignment = Alignment(
                        horizontal="center",
                        vertical="center"
                    )
            
            week_number = selected_date.isocalendar()[1]
            week_beginning_text = f"Week beginning {selected_date.strftime('%d %b')}\nWeek {week_number}"
            
            for cell in ["O2", "O33", "O65"]:
                template_sheet[cell].value = week_beginning_text
                template_sheet[cell].alignment = Alignment(
                    horizontal="center",
                    vertical="center",
                    wrap_text=True
                )
            
            output = BytesIO()
            workbook.save(output)
            output.seek(0)
            
            return output, pivot_df
        
        except Exception as e:
            logger.error(f"Failed to generate Excel report: {str(e)}")
            raise

class VenueReportGenerator:
    @staticmethod
    def generate_venue_usage_report(filtered_df: pd.DataFrame, start_date: datetime.date) -> BytesIO:
        """Generate Word document summarizing venue usage."""
        doc = Document()
        section = doc.sections[0]
        section.orientation = 1  # Landscape
        
        # Set page dimensions efficiently
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height
        
        # Add header
        title = doc.add_heading('Venue Usage Report', level=1)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph(f'Week Beginning: {start_date.strftime("%d %b %Y")}', style='Normal')
        
        # Define colors once
        DAY_COLORS = {
            "Sunday": "D3D3D3",     # light grey
            "Monday": "FFFFFF",     # white
            "Tuesday": "D3D3D3",
            "Wednesday": "FFFFFF",
            "Thursday": "D3D3D3",
            "Friday": "FFFFFF",
            "Saturday": "D3D3D3"
        }
        
        # Get unique venues once
        venues = sorted([str(v) for v in filtered_df['Venue'].dropna().unique()])
        page_capacity = 5
        
        for i in range(0, len(venues), page_capacity):
            if i > 0:
                doc.add_page_break()
            
            venue_subset = venues[i:i+page_capacity]
            for venue in venue_subset:
                venue_data = filtered_df[
                    filtered_df['Venue'].apply(lambda x: str(x)) == venue
                ].sort_values(by=['Date', 'Start_Time'])
                
                # Add venue header
                venue_heading = doc.add_heading(f'📍 {venue}', level=2)
                venue_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Create table template once
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Set headers once
                hdr_cells = table.rows[0].cells
                for cell, text in zip(hdr_cells, ['Date', 'Time', 'Training Group', 'Sport']):
                    cell.text = text
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    VenueReportGenerator._set_cell_background(cell, "ADD8E6")
                
                # Add rows efficiently
                for _, row in venue_data.iterrows():
                    cells = table.add_row().cells
                    
                    date_str = row['Date'].strftime('%A %d %b %Y')
                    time_str = f"{row['Start_Time']} - {row['Finish_Time']}"
                    
                    cells[0].text = date_str
                    cells[1].text = time_str
                    cells[2].text = str(row['Training_Group'])
                    cells[3].text = str(row['Sport'])
                    
                    day_name = row['Date'].strftime('%A')
                    color = DAY_COLORS.get(day_name, "FFFFFF")
                    for cell in cells:
                        VenueReportGenerator._set_cell_background(cell, color)
        
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output
    
    @staticmethod
    def _set_cell_background(cell, color):
        """Helper method to set cell background color."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

class TrainingApp:
    def __init__(self):
        self.excel_generator = TrainingReportGenerator()
        self.venue_generator = VenueReportGenerator()
        
        # Initialize session state once
        if "generated" not in st.session_state:
            st.session_state.generated = False
            st.session_state.excel_file = None
            st.session_state.pivot_df = None
            st.session_state.filtered_data = None
            st.session_state.venue_file = None
            
    def run(self):
        """Main application runner."""
        TrainingReportGenerator.setup_page()
        st.title("Operations - Weekly Training Plan App")
        st.markdown("Generate Training Calendar and Venue Usage reports for any week from 1st January 2025.")
        
        selected_date = st.date_input(
            "Select a starting date (make sure to choose a SUNDAY!)",
            value=datetime.now().date(),
            min_value=datetime(2025, 1, 1).date()
        )
        
        if st.button("Generate Reports"):
            try:
                excel_file, pivot_df = self.excel_generator.generate_excel_report(selected_date)
                
                # Filter data for venue report
                df = self.excel_generator.fetch_training_data()
                df['Date'] = pd.to_datetime(df['Date'], errors='coerce', dayfirst=True).dt.date
                filtered_data = df[
                    (df['Date'] >= selected_date) & 
                    (df['Date'] <= selected_date + timedelta(days=6))
                ].copy()
                
                venue_file = self.venue_generator.generate_venue_usage_report(filtered_data, selected_date)
                
                st.session_state.excel_file = excel_file
                st.session_state.pivot_df = pivot_df
                st.session_state.filtered_data = filtered_data
                st.session_state.venue_file = venue_file
                st.session_state.generated = True
                
            except Exception as e:
                logger.error(f"Error generating reports: {str(e)}")
                st.error(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    app = TrainingApp()
    app.run()