import streamlit as st
from io import BytesIO, StringIO
import requests
import pandas as pd
from datetime import datetime, timedelta, timezone
from openpyxl import load_workbook
from openpyxl.styles import Alignment
import shutil
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# ---- Page Configuration ----
st.set_page_config(
    page_title="Operations - Weekly Training Plan",
    layout="wide",  # Use wide layout
    initial_sidebar_state="expanded"
)

# ---- Custom CSS to hide default Streamlit elements and reduce top spacing ----
hide_streamlit_style = """
<style>
/* Hide the default hamburger menu and footer */
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}
</style>
"""
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Function to load and preprocess data
def load_data(selected_date):
    """Fetches and processes data from Smartabase API."""
    session = requests.Session()
    session.auth = ("kenneth.mcmillan", "Quango76")
    response = session.get("https://aspire.smartabase.com/aspireacademy/live?report=PYTHON5_TRAINING_PLAN&updategroup=true")
    response.raise_for_status()
    data = pd.read_html(StringIO(response.text))[0]
    df = data.drop(columns=['About'], errors='ignore').drop_duplicates()
    df.columns = df.columns.str.replace(' ', '_')
    
    # Convert timestamps
    def convert_to_time(timestamp_ms, offset_hours=11):
        try:
            if pd.notnull(timestamp_ms):
                timestamp_s = float(timestamp_ms) / 1000
                return (datetime.fromtimestamp(timestamp_s, tz=timezone.utc) - timedelta(hours=offset_hours)).strftime('%H:%M')
        except (ValueError, TypeError):
            return None
    
    df['Start_Time'] = pd.to_numeric(df['Start_Time'], errors='coerce').apply(lambda x: convert_to_time(x))
    df['Finish_Time'] = pd.to_numeric(df['Finish_Time'], errors='coerce').apply(lambda x: convert_to_time(x))
    df = df[df['Sport'].notna() & (df['Sport'].str.strip() != '')]
    df = df[df['Venue'] != 'AASMC']
    df = df[df['Sport'] != 'Generic Athlete']
    df = df[df['Training_Group'] != 'Practice']
    df['Date'] = pd.to_datetime(df['Date'], errors='coerce', dayfirst=True).dt.date
    
    # Filter based on selected date
    start_date = selected_date
    end_date = start_date + timedelta(days=6)
    filtered_df = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
    filtered_df.loc[:, 'AM/PM'] = pd.Categorical(filtered_df['AM/PM'], categories=['AM', 'PM'], ordered=True)
    filtered_df = filtered_df.dropna(subset=['Sport']).sort_values(by=['Date', 'Sport', 'Coach', 'AM/PM'])
    return filtered_df

# Function to generate the Excel report
def generate_excel(selected_date):
    """Generates a basic Excel file placeholder."""
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df = pd.DataFrame({'Date': [selected_date.strftime('%Y-%m-%d')], 'Message': ['Excel report generated']})
        df.to_excel(writer, index=False, sheet_name='Report')
    output.seek(0)
    return output, df

# Function to generate a nicely formatted Word document
def generate_venue_usage_report(filtered_df, start_date):
    """Generates a well-formatted Word document summarizing venue usage by day and AM/PM."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = 1  # Set to landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    title = doc.add_heading('Venue Usage Report', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f'Week Beginning: {start_date.strftime("%d %b %Y")}', style='Normal')
    
    venues = list(filtered_df['Venue'].unique())
    page_capacity = 2  # Maximum venues per page
    
    for i in range(0, len(venues), page_capacity):
        if i > 0:
            doc.add_page_break()
        
        venue_subset = venues[i:i+page_capacity]
        for venue in venue_subset:
            venue_data = filtered_df[filtered_df['Venue'] == venue].sort_values(by=['Date', 'Start_Time'])
            venue_heading = doc.add_heading(f'📍 {venue}', level=2)
            venue_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Date & Time'
            hdr_cells[1].text = 'Training Group'
            hdr_cells[2].text = 'Sport'
            
            for _, row in venue_data.iterrows():
                row_cells = table.add_row().cells
                date_time = f"{row['Date'].strftime('%A %d %b %Y')}\n{row['Start_Time']} - {row['Finish_Time']}"
                row_cells[0].text = date_time
                row_cells[1].text = row['Training_Group']
                row_cells[2].text = row['Sport']
            
            for cell in table.rows[0].cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# Streamlit App
st.title("Operations - Weekly Training Plan App")
st.markdown("Generate an Excel and Venue Usage report for any week (past or future).")

# Date input
selected_date = st.date_input("Select a starting Sunday (make sure to choose a Sunday)", value=datetime.now().date())

# Button to generate the reports
if st.button("Generate Reports"):
    try:
        # Load and process data
        filtered_df = load_data(selected_date)
        
        # Generate the Excel report
        excel_file, pivot_df = generate_excel(selected_date)

        # Display the pivot DataFrame for debugging
        st.markdown("### DataFrame for checking data")
        st.dataframe(pivot_df)

        # Provide download button for the generated Excel report
        st.download_button(
            label="Download Excel Report",
            data=excel_file,
            file_name=f"Training_Report_{selected_date.strftime('%d%b%Y')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        
        # Generate and provide download for the Venue Usage Report
        venue_file = generate_venue_usage_report(filtered_df, selected_date)
        st.download_button(
            label="Download Venue Usage Report",
            data=venue_file,
            file_name=f"Venue_Usage_Report_{selected_date.strftime('%d%b%Y')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.error(f"An error occurred: {e}")
