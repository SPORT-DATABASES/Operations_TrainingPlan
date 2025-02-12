# ---------------------------
# IMPORTS
# ---------------------------
# import streamlit as st  # <-- commented out Streamlit parts
from io import BytesIO, StringIO
import requests
import pandas as pd
from datetime import datetime, timedelta, timezone
from openpyxl import load_workbook
from openpyxl.styles import Alignment
import shutil
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn  # needed for cell shading

# ---------------------------
# (Optional) Streamlit page configuration -- commented out for debugging
# ---------------------------
# st.set_page_config(page_title="Operations - Weekly Training Plan", layout="wide", initial_sidebar_state="expanded")
# hide_streamlit_style = """
# <style>
# /* Hide the default hamburger menu and footer */
# #MainMenu {visibility: hidden;}
# footer {visibility: hidden;}
# </style>
# """
# st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# ---------------------------
# SET A FIXED STARTING DATE
# ---------------------------
# Instead of using st.date_input, just set a date:
selected_date = datetime(2025, 2, 16).date()  # Change as needed
print("Selected starting date:", selected_date.strftime("%Y-%m-%d"))

# ---------------------------
# STEP 1: Open the Excel template and copy it to an output file
# ---------------------------
template_path = "Excel_template.xlsx"  # Ensure this file is in the working directory
output_filename = f"Training_Report_{selected_date.strftime('%d%b%Y')}.xlsx"
shutil.copy(template_path, output_filename)
workbook = load_workbook(output_filename)
template_sheet = workbook["Template"]

start_date = selected_date
end_date = start_date + timedelta(days=6)
print(f"Selected Date Range: {start_date.strftime('%a %d %b %Y')} to {end_date.strftime('%a %d %b %Y')}")

# ---------------------------
# STEP 2: Retrieve the training data from the URL
# ---------------------------
session = requests.Session()
session.auth = ("sb_sap.etl", "A1s2p3!re")  # Adjust credentials if needed
response = session.get("https://aspire.smartabase.com/aspireacademy/live?report=PYTHON6_TRAINING_PLAN&updategroup=true")
response.raise_for_status()  # Stop if an HTTP error occurred
data = pd.read_html(StringIO(response.text))[0]
df = data.drop(columns=['About'], errors='ignore').drop_duplicates()
# Convert column names to strings and replace spaces with underscores
df.columns = df.columns.astype(str).str.replace(' ', '_')

# ---------------------------
# STEP 3: Convert timestamps (Start_Time and Finish_Time)
# ---------------------------
# Inline conversion: convert the numeric timestamp (in ms) to '%H:%M'
df['Start_Time'] = pd.to_numeric(df['Start_Time'], errors='coerce').apply(
    lambda x: (datetime.fromtimestamp(x/1000, tz=timezone.utc) - timedelta(hours=11)).strftime('%H:%M')
    if pd.notnull(x) else None
)
df['Finish_Time'] = pd.to_numeric(df['Finish_Time'], errors='coerce').apply(
    lambda x: (datetime.fromtimestamp(x/1000, tz=timezone.utc) - timedelta(hours=11)).strftime('%H:%M')
    if pd.notnull(x) else None
)

# ---------------------------
# STEP 4: Filter the data
# ---------------------------
df = df[df['Sport'].notna() & (df['Sport'].astype(str).str.strip() != '')]
df = df[df['Venue'] != 'AASMC']
df = df[df['Sport'] != 'Generic_Athlete']
df = df[df['Training_Group'] != 'Practice']

# Convert the 'Date' column to Python date objects (dayfirst=True)
df['Date'] = pd.to_datetime(df['Date'], errors='coerce', dayfirst=True).dt.date
filtered_df = df[(df['Date'] >= start_date) & (df['Date'] <= end_date)]
print("Filtered data (first 5 rows):")
print(filtered_df.head())

# ---------------------------
# STEP 5: Ensure key columns are strings and add a combined "Day_AM/PM" column
# ---------------------------
filtered_df['AM/PM'] = filtered_df['AM/PM'].fillna('').astype(str)
filtered_df['Session_Type'] = filtered_df['Session_Type'].fillna('').astype(str)

# Create the combined column (e.g. "Sunday AM")
filtered_df['Day_AM/PM'] = filtered_df.apply(
    lambda row: row['Date'].strftime('%A') + " " + str(row['AM/PM'])
    if pd.notnull(row['Date']) and pd.notnull(row['AM/PM']) else '',
    axis=1
)

# Sort the dataframe by Date, Sport, Coach, and AM/PM
filtered_df = filtered_df.dropna(subset=['Sport']).sort_values(by=['Date', 'Sport', 'Coach', 'AM/PM'])
print("After adding Day_AM/PM (first 5 rows):")
print(filtered_df.head())

# ---------------------------
# STEP 6: Group the data and build the pivot table
# ---------------------------
# Instead of using a helper function, we perform the grouping inline.
# We will iterate over groups defined by ['Sport', 'Training_Group', 'Day_AM/PM', 'Session_Type'].
grouped_rows = []
group_columns = ['Sport', 'Training_Group', 'Day_AM/PM', 'Session_Type']
for name, group in filtered_df.groupby(group_columns):
    sport, training_group, day_am_pm, session_type = name
    # If any row in the group is a "Training Camp", use that immediately.
    if (group['Session_Type'] == "Training Camp").any():
        session_str = "TRAINING CAMP"
    else:
        pairs = []
        for idx, row in group.iterrows():
            type_value = str(row['Session_Type']) if pd.notnull(row['Session_Type']) else ''
            venue = str(row['Venue']) if pd.notnull(row['Venue']) else ''
            start_time = str(row['Start_Time']) if pd.notnull(row['Start_Time']) else ''
            finish_time = str(row['Finish_Time']) if pd.notnull(row['Finish_Time']) else ''
            time_str = f"{start_time}-{finish_time}" if (start_time or finish_time) else ''
            if type_value == "Competition":
                formatted_entry = f"Competition\n{venue}\n{time_str}".strip()
            else:
                formatted_entry = f"{venue}\n{time_str}".strip()
            if venue or time_str or type_value:
                pairs.append((start_time, formatted_entry))
        # Define an inline parse function for sorting by time
        def parse_time(x):
            try:
                return datetime.strptime(x, '%H:%M')
            except Exception:
                return datetime.min
        pairs.sort(key=lambda x: parse_time(x[0]))
        sorted_sessions = [pair[1] for pair in pairs]
        session_str = '\n'.join([s for s in sorted_sessions if s])
    grouped_rows.append({
        'Sport': sport,
        'Training_Group': training_group,
        'Day_AM/PM': day_am_pm,
        'Session_Type': session_type,
        'Session': session_str
    })
grouped_df = pd.DataFrame(grouped_rows)

# Now pivot so that rows are indexed by Sport and Training_Group and columns are Day_AM/PM values.
pivot_df = pd.pivot_table(
    grouped_df,
    values='Session',
    index=['Sport', 'Training_Group'],
    columns=['Day_AM/PM'],
    aggfunc='first',
    fill_value=' '
).reset_index()

# Ensure that every expected Day_AM/PM column exists.
day_order = [f"{day} {time}" for day in 
             ['Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
             for time in ['AM', 'PM']]
for col in ['Sport', 'Training_Group'] + day_order:
    if col not in pivot_df.columns:
        pivot_df[col] = ' '
pivot_df = pivot_df[['Sport', 'Training_Group'] + day_order]

print("Pivot DataFrame:")
print(pivot_df)

# ---------------------------
# STEP 7: Paste the pivot data into the Excel template
# ---------------------------
# Define the list of rows to paste. (These indicate which sport/training group goes in which cell.)
rows_to_paste = [
    {"sport": "Development", "training_group": "Development 1", "start_cell": "C6"},
    {"sport": "Development", "training_group": "Development 2", "start_cell": "C8"},
    {"sport": "Development", "training_group": "Development 3", "start_cell": "C10"},
    {"sport": "Endurance", "training_group": "Endurance_Senior", "start_cell": "C12"},
    {"sport": "Jumps", "training_group": "Jumps_PV", "start_cell": "C14"},
    {"sport": "Jumps", "training_group": "Jumps_Martin Bercel", "start_cell": "C16"},
    {"sport": "Jumps", "training_group": "Jumps_Ross Jeffs", "start_cell": "C18"},
    {"sport": "Jumps", "training_group": "Jumps_ElWalid", "start_cell": "C20"},
    {"sport": "Sprints", "training_group": "Sprints_Lee", "start_cell": "C22"},
    {"sport": "Sprints", "training_group": "Sprints_Hamdi", "start_cell": "C24"},
    {"sport": "Throws", "training_group": "Senior Performance Throws", "start_cell": "C26"},
    {"sport": "Squash", "training_group": "Squash", "start_cell": "C37"},
    {"sport": "Table Tennis", "training_group": "Table Tennis", "start_cell": "C39"},
    {"sport": "Fencing", "training_group": "Fencing", "start_cell": "C41"},
    {"sport": "Swimming", "training_group": "Swimming", "start_cell": "C43"},
    {"sport": "Padel", "training_group": "Padel", "start_cell": "C45"},
    {"sport": "Pre Academy Padel", "training_group": "Explorers", "start_cell": "C48"},
    {"sport": "Pre Academy Padel", "training_group": "Explorers+", "start_cell": "C49"},
    {"sport": "Pre Academy Padel", "training_group": "Starters", "start_cell": "C50"},
    {"sport": "Pre Academy", "training_group": "Pre Academy Fencing", "start_cell": "C51"},
    {"sport": "Pre Academy", "training_group": "Pre Academy Squash Girls", "start_cell": "C53"},
    {"sport": "Pre Academy", "training_group": "Pre Academy Athletics", "start_cell": "C55"},
    {"sport": "Girls Programe", "training_group": "Kids", "start_cell": "C58"},
    {"sport": "Girls Programe", "training_group": "Mini Cadet_U14", "start_cell": "C59"},
    {"sport": "Girls Programe", "training_group": "Cadet_U16", "start_cell": "C60"},
    {"sport": "Girls Programe", "training_group": "Youth_U18", "start_cell": "C61"},
    {"sport": "Sprints", "training_group": "Sprints_Short", "start_cell": "C69"},
    {"sport": "Sprints", "training_group": "Sprints_Long", "start_cell": "C71"},
    {"sport": "Jumps", "training_group": "Jumps_QAF", "start_cell": "C73"},
    {"sport": "Throws", "training_group": "Discus_QAF", "start_cell": "C75"},
    {"sport": "Throws", "training_group": "Hammer_QAF", "start_cell": "C77"},
    {"sport": "Throws", "training_group": "Javelin_QAF", "start_cell": "C79"}
]

for row in rows_to_paste:
    sport_val = row["sport"]
    training_group_val = row["training_group"]
    start_cell = row["start_cell"]
    # Filter the pivot_df for the matching sport and training group:
    filtered_row = pivot_df[(pivot_df['Sport'] == sport_val) & (pivot_df['Training_Group'] == training_group_val)]
    if not filtered_row.empty:
        # Get values from the pivot (skip the first two columns)
        values_to_paste = filtered_row.iloc[0, 2:].tolist()
        col_letter = start_cell[0]
        row_num = int(start_cell[1:])
        start_col_idx = ord(col_letter.upper()) - ord("A") + 1
        for col_idx, value in enumerate(values_to_paste, start=start_col_idx):
            cell = template_sheet.cell(row=row_num, column=col_idx, value=value)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

# ---------------------------
# STEP 8: Set date cells in the Excel template
# ---------------------------
date_cells_groups = [
    ['C4', 'E4', 'G4', 'I4', 'K4', 'M4', 'O4'],
    ['C35', 'E35', 'G35', 'I35', 'K35', 'M35', 'O35'],
    ['C67', 'E67', 'G67', 'I67', 'K67', 'M67', 'O67'],
]
for day_offset, cell_group in enumerate(zip(*date_cells_groups)):
    date_value = (start_date + timedelta(days=day_offset)).strftime('%a %d %b %Y')
    for cell in cell_group:
        template_sheet[cell].value = date_value
        template_sheet[cell].alignment = Alignment(horizontal="center", vertical="center")

week_number = start_date.isocalendar()[1]
week_beginning_text = f"Week beginning {start_date.strftime('%d %b')}\nWeek {week_number}"
for cell in ["O2", "O33", "O65"]:
    template_sheet[cell].value = week_beginning_text
    template_sheet[cell].alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

# Save the modified Excel workbook to a BytesIO buffer and/or to disk for debugging.
excel_output = BytesIO()
workbook.save(excel_output)
excel_output.seek(0)
with open(f"Debug_Training_Report_{selected_date.strftime('%d%b%Y')}.xlsx", "wb") as f:
    f.write(excel_output.getbuffer())
print("Excel report generated and saved.")

# ---------------------------
# STEP 9: Generate the Word venue usage report inline
# ---------------------------
doc = Document()
section = doc.sections[0]
# Set landscape orientation:
section.orientation = 1  
new_width, new_height = section.page_height, section.page_width
section.page_width = new_width
section.page_height = new_height

heading = doc.add_heading('Venue Usage Report', level=1)
heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph(f'Week Beginning: {start_date.strftime("%d %b %Y")}', style='Normal')

# Define colors for days:
day_colors = {
    "Sunday": "D3D3D3",     # light grey
    "Monday": "FFFFFF",     # white
    "Tuesday": "D3D3D3",
    "Wednesday": "FFFFFF",
    "Thursday": "D3D3D3",
    "Friday": "FFFFFF",
    "Saturday": "D3D3D3"
}

# Get the list of venues (alphabetically)
venues = sorted([str(v) for v in filtered_df['Venue'].dropna().unique()])
page_capacity = 5  # maximum venues per page

for i in range(0, len(venues), page_capacity):
    if i > 0:
        doc.add_page_break()
    venue_subset = venues[i:i+page_capacity]
    for venue in venue_subset:
        # For each venue, filter the data and sort by Date and Start_Time
        venue_data = filtered_df[filtered_df['Venue'].apply(lambda x: str(x)) == venue].sort_values(by=['Date', 'Start_Time'])
        venue_heading = doc.add_heading(f'📍 {venue}', level=2)
        venue_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Create a table with 4 columns: Date, Time, Training Group, Sport
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Date'
        hdr_cells[1].text = 'Time'
        hdr_cells[2].text = 'Training Group'
        hdr_cells[3].text = 'Sport'
        # Bold and center the header cells; set a light blue background
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), "ADD8E6")
            tcPr.append(shd)
        
        # Add one row per record for this venue
        for idx, row in venue_data.iterrows():
            row_cells = table.add_row().cells
            date_str = row['Date'].strftime('%A %d %b %Y')
            time_str = f"{row['Start_Time']} - {row['Finish_Time']}"
            row_cells[0].text = date_str
            row_cells[1].text = time_str
            row_cells[2].text = str(row['Training_Group'])
            row_cells[3].text = str(row['Sport'])

            # Set background color based on day-of-week
            day_name = row['Date'].strftime('%A')
            color = day_colors.get(day_name, "FFFFFF")
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), color)
                tcPr.append(shd)

word_output = BytesIO()
doc.save(word_output)
word_output.seek(0)
with open(f"Debug_Venue_Usage_Report_{selected_date.strftime('%d%b%Y')}.docx", "wb") as f:
    f.write(word_output.getbuffer())
print("Word venue usage report generated and saved.")
